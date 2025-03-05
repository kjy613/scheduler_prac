import schedule
import time
import smtplib
from docx import Document
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.base import MIMEBase
from email import encoders

# 템플릿 파일 경로
TEMPLATE_PATH = r"c:\scheduler\template.docx"
OUTPUT_PATH = r"c:\scheduler\output.docx"
EMAIL_SENDER = "kkosoljyli@gmail.com"
EMAIL_PW = "---- ---- ---- ----" # 구글 앱 비밀번호
EMAIL_RECEIVER = "jyeon@castis.com"

def create_word_file():
    doc = Document(TEMPLATE_PATH)
    today_date = datetime.today().strftime("%Y-%m-%d")
    current_time = datetime.now().strftime("%H:%M")

    for para in doc.paragraphs:
        if "{date}" in para.text:
            para.text = para.text.replace("{date}", today_date)

    for para in doc.paragraphs:
        if "{time}" in para.text:
            para.text = para.text.replace("{time}", current_time)

    doc.save(OUTPUT_PATH)
    print(f"파일 생성 완료: {OUTPUT_PATH}")
          
def send_email():
        msg = MIMEMultipart()
        msg["From"] = EMAIL_SENDER
        msg["To"] = EMAIL_RECEIVER
        msg["Subject"] = " 스케줄러 연습(자동 생성 문서) "

        with open(OUTPUT_PATH, "rb") as attachment:
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header("Content-Disposition", f"attachment; filename={OUTPUT_PATH}")
            msg.attach(part)

        with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
            server.login(EMAIL_SENDER, EMAIL_PW)
            server.sendmail(EMAIL_SENDER, EMAIL_RECEIVER, msg.as_string())

        print(f" 이메일 전송 완료: {EMAIL_RECEIVER}")

def run_task():
    # print("작업 시작")
    create_word_file()
    send_email()
    # print("작업 완료")


# 매일, 9시에 실행
schedule.every().days.at("09:00").do(run_task)
# run_task()로 하면 함수 호출이 아니라 함수의 결과로 사용됨. ()붙으면 함수 즉시실행,반환값 do()에 전달하게됨.

print("자동 스케줄러 실행 중...")
while True:
    schedule.run_pending()
    time.sleep(60)
